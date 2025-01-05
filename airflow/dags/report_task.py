from airflow import DAG
from airflow.operators.python import PythonOperator
from datetime import datetime
from pymongo import MongoClient
from docx import Document
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from google.auth.transport.requests import Request
from google.oauth2.service_account import Credentials
from googleapiclient.errors import HttpError
import os
import re
from time import sleep

SCOPES = ['https://www.googleapis.com/auth/drive.file']

# Function to connect to MongoDB
def connect_to_mongo():
    uri = "mongodb+srv://simchab667:GesJiqzTeqyHsdp5@cluster0.yhfpz1v.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0"
    client = MongoClient(uri)
    return client

# Function to authenticate with Google Drive using Service Account
def authenticate_google_drive():
    service_account_file = '/opt/airflow/dags/service_account.json'  
    credentials = Credentials.from_service_account_file(
        service_account_file, scopes=SCOPES)
    return build('drive', 'v3', credentials=credentials)

# Function to upload a file to Google Drive
def upload_to_drive(filename):
    service = authenticate_google_drive()
    file_metadata = {
        'name': os.path.basename(filename),
        'parents': ['1MCoZ9GJAkIei61I_m2rRhCW44bLbbMpo']  
    }
    media = MediaFileUpload(
        filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

    for attempt in range(5): 
        try:
            file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
            print(f"הקובץ הועלה בהצלחה ל-Google Drive. File ID: {file.get('id')}")
            return file.get('id')
        except HttpError as error:
            print(f"שגיאה ב-HTTP: {error}")
            if attempt < 4:
                print("ממתינים 5 שניות לפני ניסיון חוזר...")
                sleep(5)  
            else:
                raise

# Function to delete a file from Google Drive
def delete_file_from_drive(file_id):
    service = authenticate_google_drive()
    try:
        service.files().delete(fileId=file_id).execute()
        print(f"קובץ עם ID {file_id} נמחק בהצלחה.")
    except HttpError as error:
        print(f"שגיאה במחיקת קובץ מגוגל דרייב: {error}")

# Function to check if the description contains a link to Google Drive
def extract_drive_file_id(link):
    match = re.search(r"https://drive.google.com/file/d/([a-zA-Z0-9_-]+)/view", link)
    return match.group(1) if match else None

# Function to create a share link for the file
def get_drive_share_link(file_id, service):
    permission = {
        'type': 'anyone',  # Public share link
        'role': 'reader'  # Read only
    }
    service.permissions().create(fileId=file_id, body=permission).execute()
    share_link = f"https://drive.google.com/file/d/{file_id}/view?usp=sharing"
    return share_link

# Function to create separate reports for completed tasks and upload them to Google Drive
def get_completed_task_details():
    client = connect_to_mongo()
    db = client.get_database('MyDB')
    tasks_collection = db.get_collection('task')
    sensors_collection = db.get_collection('sensor')

    all_tasks = tasks_collection.find()

    for task in all_tasks:
        link = task.get('report_link', '')  
        drive_file_id = extract_drive_file_id(link)  

        if task['status'] == "הושלם":
            if not drive_file_id:  
               # Create a new report
                document = Document()
                document.add_heading(f'דוח עבור משימה: {task["name"]}', level=1)
                document.add_paragraph(f"תיאור: {task['description']}")
                document.add_paragraph(f"בעלים: {task['owner']}")
                document.add_paragraph(f"סטטוס: {task['status']}")
                document.add_paragraph(f"צוות: {task['team']}")
                document.add_paragraph(f"התחלה: {datetime_to_str(task['start_date'])}")
                document.add_paragraph(f"סיום: {datetime_to_str(task['end_date'])}")

              # Adding sensor information
                if 'sensors' in task:
                    document.add_heading("סנסורים:", level=2)
                    for sensor_id in task['sensors']:
                        sensor = sensors_collection.find_one({'id': sensor_id})
                        if sensor:
                            document.add_paragraph(f"סנסור: {sensor['name']}", style='List Bullet')
                            document.add_paragraph(f"סטטוס: {sensor['status']}")

               # Temporarily save before uploading
                report_filename = f"{task['name']}.docx"
                temp_report_path = f"/tmp/{report_filename}"
                document.save(temp_report_path)

                # Upload to Google Drive
                file_id = upload_to_drive(temp_report_path)
                service = authenticate_google_drive()
                share_link = get_drive_share_link(file_id, service)

                # Update the link field in the collection
                tasks_collection.update_one(
                    {'_id': task['_id']},
                    {'$set': {'report_link': share_link}}
                )
                print(f"דוח עבור המשימה {task['name']} הועלה וקישור השיתוף עודכן.")
        else:
            if drive_file_id: 
                delete_file_from_drive(drive_file_id)  # Delete the file on the drive
                tasks_collection.update_one(
                    {'_id': task['_id']},
                    {'$set': {'report_link': ''}}  # Remove the link from the
                )
                print(f"קישור השיתוף עבור המשימה {task['name']} הוסר.")

# Function to convert datetime objects to a string in date and time format
def datetime_to_str(dt):
    if isinstance(dt, datetime):
        return dt.strftime('%Y-%m-%d %H:%M:%S')
    return None

# Setting up the DAG
default_args = {
    'owner': 'Simcha',
    'depends_on_past': False,
    'start_date': datetime(2024, 12, 5), 
    'retries': 1,
}

dag = DAG(
    'completed_task_report',
    default_args=default_args,
    description='DAG to generate a report for completed tasks and upload it to Google Drive',
    schedule_interval="0 0 * * *",  # Run every day at 00:00 AM
    catchup=False,  # Undo retroactive task completion
)

# Task to create the report and upload it to Drive
generate_and_upload_report_task = PythonOperator(
    task_id='generate_and_upload_report',
    python_callable=get_completed_task_details,
    dag=dag,
)